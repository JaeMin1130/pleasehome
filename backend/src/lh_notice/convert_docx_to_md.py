import os
import glob
import docx

def convert_docx_to_md():
    guide_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "docs", "api-guide")
    docx_files = glob.glob(os.path.join(guide_dir, "*.docx"))
    
    if not docx_files:
        print("변환할 docx 파일이 없습니다.")
        return

    for docx_path in docx_files:
        md_path = docx_path.rsplit(".", 1)[0] + ".md"
        
        # 이미 변환된 md 파일이 있으면 건너뛰기
        if os.path.exists(md_path):
            print(f"이미 변환됨: {os.path.basename(md_path)}")
            continue
            
        print(f"변환 중: {os.path.basename(docx_path)}")
        try:
            doc = docx.Document(docx_path)
            lines = []
            
            # 본문 문단 추출
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    lines.append(text)
                    
            lines.append("\n--- [데이터 테이블] ---\n")
            
            # 표 데이터 추출 (파라미터 등)
            for t in doc.tables:
                for row_idx, row in enumerate(t.rows):
                    row_data = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    lines.append(" | ".join(row_data))
                    # 표 헤더 아래에 마크다운 구분선 추가
                    if row_idx == 0:
                        lines.append(" | ".join(["---"] * len(row.cells)))
                lines.append("") # 표 사이 여백
                
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
                
            print(f" -> 성공: {os.path.basename(md_path)}")
        except Exception as e:
            print(f" -> 실패: {os.path.basename(docx_path)} (사유: {e})")

if __name__ == "__main__":
    convert_docx_to_md()
